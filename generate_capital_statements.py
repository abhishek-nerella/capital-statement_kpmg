"""
Capital Analysis Statement generator.
Usage: python generate_capital_statements.py --input investors.xlsx --output ./output/
"""

import argparse
import os
import re
import sys

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


# ── formatting helpers ────────────────────────────────────────────────────────

def fmt_usd(value) -> str:
    """Format a numeric value as $1,234,567.89; negatives as ($1,234,567.89)."""
    try:
        v = float(value)
    except (TypeError, ValueError):
        return "$0.00"
    if v < 0:
        return f"(${abs(v):,.2f})"   # Change 4: bracket notation for negatives
    return f"${v:,.2f}"


def fmt_ratio(value) -> str:
    """Format TEV_RATIO as 0.72x"""
    try:
        return f"{float(value):.2f}x"
    except (TypeError, ValueError):
        return "0.00x"


def fmt_date(value) -> str:
    """Return a human-readable date string from a pandas Timestamp or string."""
    if pd.isnull(value):
        return ""
    if hasattr(value, "strftime"):
        return value.strftime("%B %d, %Y")
    return str(value)


# ── label normalisation helpers ───────────────────────────────────────────────

def _norm_label(label: str) -> str:
    """Change 5: insert slash before parenthetical alternative in appreciation/gain labels."""
    label = re.sub(r"(?i)appreciation \(depreciation\)", "appreciation/(depreciation)", label)
    label = re.sub(r"(?i)gain \(loss\)", "gain/(loss)", label)
    return label


def _strip_brackets(label: str) -> str:
    """Change 3: remove all ( ) characters from a label string (Sections 2 & 3 only)."""
    return label.replace("(", "").replace(")", "")


# Change 6: columns that may be blank for transfer-case investors
NULL_COERCE_COLS = (
    "INCEPTION_TO_DATE_CONTRIBUTION",
    "INCEPTION_TO_DATE_DISTRIBUTION",
    "OPENING_YTD_NAV",
    "YTD_CONTRIBUTION",
    "YTD_DISTRIBUTION",
)


def coerce_transfer_case_nulls(row: pd.Series) -> pd.Series:
    """Change 6: null→0 coercion for blank transfer-case columns. Returns a copy.

    Shared by every builder (Word + PDF, PE + HF) so a blank ITD contribution
    (a transfer-case investor) never crashes generation for one format but not
    the other.
    """
    row = row.copy()
    investor_name = str(row.get("INVESTOR_NAME", "UNKNOWN"))
    for col in NULL_COERCE_COLS:
        val = row.get(col)
        if pd.isna(val):
            print(f"WARN: {investor_name} — {col} was null, defaulted to 0")
            row[col] = 0.0
    return row


# ── docx helpers ──────────────────────────────────────────────────────────────

def _set_run_font(run, bold=False, underline=False, size_pt=11):
    run.bold = bold
    run.underline = underline
    run.font.size = Pt(size_pt)


def add_bold_underline_paragraph(doc: Document, text: str) -> None:
    """Add a paragraph whose entire text is bold + underlined."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, bold=True, underline=True)


def add_two_col_row(doc: Document, label: str, value: str, indent: str = "") -> None:
    """Add a single-paragraph row with label left and value right (tab-separated)."""
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(4.5), WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(f"{indent}{label}\t{value}")
    run.font.size = Pt(11)


def add_blank(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def set_cell_border_bottom(cell) -> None:
    """Add a bottom border to a table cell (used for totals lines)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")
    tcBorders.append(bottom)
    tcPr.append(tcBorders)


def _clear_table_borders(table) -> None:
    """Remove all borders from a table so individual cell borders can be applied cleanly."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "none")
        tblBorders.append(elem)
    tblPr.append(tblBorders)


def _set_cell_double_border(cell) -> None:
    """Add top and bottom KPMG Blue (#00338D) single-line borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(existing)
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom"):
        elem = OxmlElement(f"w:{side}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "00338D")
        tcBorders.append(elem)
    tcPr.append(tcBorders)


def add_two_col_row_double_border(doc: Document, label: str, value: str, indent: str = "") -> None:
    """Two-col row where only the value (right) cell has top+bottom KPMG Blue borders."""
    tbl = doc.add_table(rows=1, cols=2)
    _clear_table_borders(tbl)
    tbl.columns[0].width = Inches(4.5)
    tbl.columns[1].width = Inches(2.0)

    row = tbl.rows[0]

    lc = row.cells[0]
    lc.text = f"{indent}{label}"
    for run in lc.paragraphs[0].runs:
        run.font.size = Pt(11)

    vc = row.cells[1]
    vc.text = value
    vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in vc.paragraphs[0].runs:
        run.font.size = Pt(11)
    _set_cell_double_border(vc)


# ── document builder ──────────────────────────────────────────────────────────

def build_document(row: pd.Series) -> Document:

    # Change 6: null→0 coercion for blank transfer-case columns
    row = coerce_transfer_case_nulls(row)

    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── "CONFIDENTIAL" top-right ─────────────────────────────────────────────
    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    conf_run = conf_para.add_run("CONFIDENTIAL")
    conf_run.bold = True
    conf_run.font.size = Pt(9)
    conf_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    # ── Header ───────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_para.add_run("Capital Analysis")
    t_run.bold = True
    t_run.font.size = Pt(16)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub_para.add_run(f"for the period ended {fmt_date(row['TO_DATE'])}")
    s_run.font.size = Pt(12)

    # Partnership / Investor info
    add_blank(doc)
    p = doc.add_paragraph()
    p.add_run("Partnership: ").bold = True
    p.add_run(str(row["PARTNERSHIP_NAME"]))

    p2 = doc.add_paragraph()
    p2.add_run("Investor: ").bold = True
    p2.add_run(str(row["INVESTOR_NAME"]))

    p3 = doc.add_paragraph()
    p3.add_run("Investor ID: ").bold = True
    p3.add_run(str(row["INVESTOR_ID"]))

    p4 = doc.add_paragraph()
    p4.add_run("Currency: ").bold = True
    p4.add_run(str(row["CURRENCY_CODE"]))

    add_blank(doc)

    # ── Section 1: Summary of Capital Account ────────────────────────────────
    add_bold_underline_paragraph(doc, "Summary of Capital Account")
    add_blank(doc)

    # Change 1: skip row if abs(value) < 0.005
    # Change 5: _norm_label fixes "appreciation (depreciation)" and "gain (loss)"

    _v = float(row["OPENING_YTD_NAV"])
    if abs(_v) >= 0.005:
        add_two_col_row(
            doc,
            _norm_label(f"Opening Capital balance as on {fmt_date(row['FROM_DATE'])}"),
            fmt_usd(_v),
        )

    _v = float(row["YTD_CONTRIBUTION"])
    if abs(_v) >= 0.005:
        add_two_col_row(doc, _norm_label("Capital contributions during the year"), fmt_usd(_v))

    _v = float(row["YTD_DISTRIBUTION"])
    if abs(_v) >= 0.005:
        add_two_col_row(doc, _norm_label("Distributions during the year"), fmt_usd(_v))

    add_blank(doc)
    p_net = doc.add_paragraph()
    p_net.add_run("Net investment activity:").font.size = Pt(11)

    net_income = float(row["INVESTMENT_INCOME"]) - float(row["INVESTMENT_EXPENSE"])
    if abs(net_income) >= 0.005:
        add_two_col_row(
            doc,
            _norm_label("Investment and other income"),
            fmt_usd(net_income),
            indent="    ",
        )

    _v = float(row["UNREALIZED_GAINS_LOSS"])
    if abs(_v) >= 0.005:
        # _norm_label converts "appreciation (depreciation)" → "appreciation/(depreciation)"
        add_two_col_row(
            doc,
            _norm_label("Net unrealized appreciation (depreciation)"),
            fmt_usd(_v),
            indent="    ",
        )

    _v = float(row["REALIZED_GAINS_LOSS"])
    if abs(_v) >= 0.005:
        # _norm_label converts "gain (loss)" → "gain/(loss)"
        add_two_col_row(
            doc,
            _norm_label("Net realized gain (loss)"),
            fmt_usd(_v),
            indent="    ",
        )

    add_blank(doc)

    _v = float(row["MANAGEMENT_FEE"])
    if abs(_v) >= 0.005:
        add_two_col_row(doc, _norm_label("Management fees for the period"), fmt_usd(_v))

    _v = float(row["INCENTIVE_FEE"])
    if abs(_v) >= 0.005:
        add_two_col_row(doc, _norm_label("Incentive fees for the period"), fmt_usd(_v))

    add_blank(doc)

    # Change 2: double border on Ending NAV (amount cell only)
    _v = float(row["CLOSING_YTD_NAV"])
    if abs(_v) >= 0.005:
        add_two_col_row_double_border(
            doc,
            _norm_label(f"Capital balance (remaining value) at {fmt_date(row['TO_DATE'])} *"),
            fmt_usd(_v),
        )

    add_blank(doc)

    # ── Section 2: Summary of Capital Commitment ─────────────────────────────
    add_bold_underline_paragraph(doc, "Summary of Capital Commitment")
    add_blank(doc)

    committed   = float(row["COMMITTED_CAPITAL"])
    contributed = float(row["INCEPTION_TO_DATE_CONTRIBUTION"])
    remaining   = committed - contributed

    # Change 3: _strip_brackets removes ( ) from labels in this section
    if abs(committed) >= 0.005:
        add_two_col_row(
            doc,
            _norm_label(_strip_brackets("Capital commitment per subscription agreement (A)")),
            fmt_usd(committed),
        )

    if abs(contributed) >= 0.005:
        add_two_col_row(
            doc,
            _norm_label(_strip_brackets("Capital contributed to date (B)")),
            fmt_usd(contributed),
        )

    # Change 2: double border on Unfunded Commitment (amount cell only)
    if abs(remaining) >= 0.005:
        add_two_col_row_double_border(
            doc,
            _norm_label(_strip_brackets("Remaining capital commitment (A-B)")),
            fmt_usd(remaining),
        )

    add_blank(doc)

    # ── Section 3: Summary of Distributions and Valuation ───────────────────
    add_bold_underline_paragraph(doc, "Summary of Distributions and Valuation")
    add_blank(doc)

    # Change 3: _strip_brackets removes ( ) from labels in this section

    _v = float(row["INCEPTION_TO_DATE_CONTRIBUTION"])
    if abs(_v) >= 0.005:
        add_two_col_row(
            doc,
            _norm_label(_strip_brackets("Total capital contributed to date")),
            fmt_usd(_v),
        )

    _v = float(row["INCEPTION_TO_DATE_DISTRIBUTION"])
    if abs(_v) >= 0.005:
        add_two_col_row(
            doc,
            _norm_label(_strip_brackets("Total distributions to date")),
            fmt_usd(_v),
        )

    _v = float(row["CLOSING_YTD_NAV"])
    if abs(_v) >= 0.005:
        add_two_col_row(
            doc,
            _norm_label(_strip_brackets(
                f"Capital balance (remaining value) at {fmt_date(row['TO_DATE'])} *"
            )),
            fmt_usd(_v),
        )

    # Change 2: double border on TEV (amount cell only)
    _v = float(row["TEV"])
    if abs(_v) >= 0.005:
        add_two_col_row_double_border(
            doc,
            _norm_label(_strip_brackets("Total Estimated Value (distributions + balance)")),
            fmt_usd(_v),
        )

    # Change 2: double border on TEV Ratio — no zero-suppression (ratio, not monetary)
    add_two_col_row_double_border(
        doc,
        _norm_label(_strip_brackets("Total Estimated Value as net multiple")),
        fmt_ratio(row["TEV_RATIO"]),
    )

    add_blank(doc)
    add_blank(doc)

    # ── Footer footnote ───────────────────────────────────────────────────────
    footnote_para = doc.add_paragraph()
    footnote_run = footnote_para.add_run(
        "* Represents remaining value. The remaining value is based upon available "
        "information and may not represent amounts which might ultimately be realized."
    )
    footnote_run.font.size = Pt(9)
    footnote_run.italic = True

    return doc


# ── main ──────────────────────────────────────────────────────────────────────

REQUIRED_COLS = {
    "FROM_DATE", "TO_DATE", "PARTNERSHIP_NAME", "INVESTOR_NAME",
    "INVESTOR_ID", "CURRENCY_CODE",
    "COMMITTED_CAPITAL", "INCEPTION_TO_DATE_CONTRIBUTION",
    "INCEPTION_TO_DATE_DISTRIBUTION", "OPENING_YTD_NAV",
    "YTD_CONTRIBUTION", "YTD_DISTRIBUTION",
    "INVESTMENT_INCOME", "INVESTMENT_EXPENSE",
    "UNREALIZED_GAINS_LOSS", "REALIZED_GAINS_LOSS",
    "MANAGEMENT_FEE", "INCENTIVE_FEE",
    "CLOSING_YTD_NAV", "TEV", "TEV_RATIO",
}

OUTPUT_COL_ORDER = [
    "FROM_DATE", "TO_DATE", "PARTNERSHIP_NAME", "INVESTOR_NAME",
    "INVESTOR_ID", "CURRENCY_CODE",
    "COMMITTED_CAPITAL", "INCEPTION_TO_DATE_CONTRIBUTION",
    "INCEPTION_TO_DATE_DISTRIBUTION", "OPENING_YTD_NAV",
    "YTD_CONTRIBUTION", "YTD_DISTRIBUTION",
    "INVESTMENT_INCOME", "INVESTMENT_EXPENSE",
    "UNREALIZED_GAINS_LOSS", "REALIZED_GAINS_LOSS",
    "MANAGEMENT_FEE", "INCENTIVE_FEE",
    "CLOSING_YTD_NAV", "TEV", "TEV_RATIO",
    "MIN_FROM_DATE", "MAX_TO_DATE", "RECORD_COUNT",
]


def build_summary_excel(df_raw: pd.DataFrame, df_latest: pd.DataFrame) -> pd.DataFrame:
    """Return a DataFrame matching the output Excel schema (one row per investor)."""
    agg = df_raw.groupby("INVESTOR_NAME").agg(
        MIN_FROM_DATE=("FROM_DATE", "min"),
        MAX_TO_DATE=("TO_DATE", "max"),
        RECORD_COUNT=("TO_DATE", "count"),
    ).reset_index()

    out = df_latest.merge(agg, on="INVESTOR_NAME", how="left")
    return out[OUTPUT_COL_ORDER]


def main():
    parser = argparse.ArgumentParser(description="Generate Capital Analysis Statements per investor.")
    parser.add_argument("--input", required=True, help="Path to the input .xlsx file")
    parser.add_argument("--output", required=True, help="Directory to write .docx files into")
    args = parser.parse_args()

    # Load data
    try:
        df = pd.read_excel(args.input)
    except Exception as exc:
        print(f"❌ Could not read input file: {exc}", file=sys.stderr)
        sys.exit(1)

    missing = REQUIRED_COLS - set(df.columns)
    if missing:
        print(f"❌ Input file is missing columns: {', '.join(sorted(missing))}", file=sys.stderr)
        sys.exit(1)

    # Parse dates for proper sorting
    df["TO_DATE"] = pd.to_datetime(df["TO_DATE"], errors="coerce")
    df["FROM_DATE"] = pd.to_datetime(df["FROM_DATE"], errors="coerce")

    # Keep only the latest TO_DATE row per investor
    df_latest = (
        df.sort_values("TO_DATE", ascending=True)
        .groupby("INVESTOR_NAME", as_index=False)
        .last()
    )

    # Create output directory
    os.makedirs(args.output, exist_ok=True)

    for _, row in df_latest.iterrows():
        investor = str(row["INVESTOR_NAME"]).strip()
        try:
            doc = build_document(row)
            safe_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in investor)
            out_path = os.path.join(args.output, f"{safe_name}_capital_statement.docx")
            doc.save(out_path)
            print(f"✅ {investor}")
        except Exception as exc:
            print(f"❌ {investor}: {exc}")

    # Write output summary Excel
    excel_path = os.path.join(args.output, "capital_statements_summary.xlsx")
    summary_df = build_summary_excel(df, df_latest)
    with pd.ExcelWriter(excel_path, engine="openpyxl", date_format="YYYY-MM-DD") as writer:
        summary_df.to_excel(writer, index=False, sheet_name="CapitalStatements")
    print(f"\n📊 Summary Excel → {excel_path}")


if __name__ == "__main__":
    main()
