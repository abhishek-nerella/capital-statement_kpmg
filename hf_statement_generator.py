"""Hedge Fund Capital Account Statement — Word (.docx) and PDF builder.

Column names from hf_pcap_engine.read_hf_pcap():
  BEG_CAP_CQ/YTD/ITD, END_CAP_CQ/YTD/ITD, BEG_UNITS_*/END_UNITS_*,
  BEG_PX/END_PX, GROSS_IRR, NET_IRR, DPI/RVPI/TVPI,
  HURDLE_AMT_ITD, LP_NET_WF, LOCKUP_MO, CUSTODIAN, etc.
"""

from __future__ import annotations

import io
import pandas as pd

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors as rl_colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from generate_capital_statements import fmt_usd, fmt_date, fmt_ratio


# ── Safe accessors ─────────────────────────────────────────────────────────────

def _g(row: pd.Series, field: str, default: float = 0.0) -> float:
    val = row.get(field, default)
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return default
    try:
        return float(val)
    except (TypeError, ValueError):
        return default


def _gs(row: pd.Series, field: str, default: str = "—") -> str:
    val = row.get(field, default)
    s = str(val).strip() if val is not None else ""
    return s if s and s not in ("nan", "None", "") else default


def _check_cq_arithmetic(row: pd.Series) -> None:
    """Change 7: verify the CQ capital roll-forward balances to within $1.00.

    Formula: BEG + contributions + income + gains − distributions − fees = END
    Raises ValueError naming the investor and the discrepancy if > $1.00.
    """
    expected = (
        _g(row, "BEG_CAP_CQ")
        + _g(row, "CONTRIB_CQ")
        + _g(row, "DRIP_CQ")
        - _g(row, "REDEMP_CQ")
        + _g(row, "XFER_IN_CQ")
        - _g(row, "XFER_OUT_CQ")
        + _g(row, "INC_CQ")
        + _g(row, "EXP_CQ")  # stored negative in PCAP — add as-is, don't double-negate
        + _g(row, "UNRLZ_CQ")
        + _g(row, "RLZD_CQ")
        - _g(row, "DIST_LP_CQ")
        - _g(row, "DIST_MGR_CQ")
        - _g(row, "INC_FEE_CQ")
        - _g(row, "TAX_RED_CQ")
    )
    actual = _g(row, "END_CAP_CQ")
    discrepancy = abs(actual - expected)
    if discrepancy > 1.00:
        investor = _gs(row, "INVESTOR_NAME")
        raise ValueError(
            f"{investor}: CQ ending capital discrepancy of ${discrepancy:,.2f} "
            f"(roll-forward expected {fmt_usd(expected)}, source END_CAP_CQ = {fmt_usd(actual)})"
        )


# ── Formatting helpers ─────────────────────────────────────────────────────────

def _fmt_pct(v) -> str:
    try:
        return f"{float(v):.2f}%"
    except Exception:
        return "—"


def _fmt_x(v) -> str:
    try:
        return f"{float(v):.2f}x"
    except Exception:
        return "—"


def _fmt_units(v) -> str:
    try:
        f = float(v)
        return f"{f:,.4f}" if f != 0.0 else "—"
    except Exception:
        return "—"


def _fmt_px(v) -> str:
    try:
        f = float(v)
        return f"${f:,.4f}" if f != 0.0 else "—"
    except Exception:
        return "—"


def _fmt_mo(v) -> str:
    try:
        mo = int(float(v))
        return f"{mo} months" if mo > 0 else "—"
    except Exception:
        return "—"


def _fmt_int(v, suffix: str = "") -> str:
    try:
        n = int(float(v))
        return f"{n}{suffix}" if n > 0 else "—"
    except Exception:
        return "—"


# ── KPMG brand constants ───────────────────────────────────────────────────────
_KPMG_BLUE     = RGBColor(0x00, 0x33, 0x8D)
_KPMG_BLUE_HEX = "#00338D"
_PACIFIC_HEX   = "#00B8F5"
_RED_HEX       = "#D73B3E"
_LIGHT_HEX     = "#ACEAFF"
_TOTAL_HEX     = "#E8F4FF"


# ── ReportLab styles (module-level, unique hf_ prefix) ────────────────────────
_hf_base = getSampleStyleSheet()
_hf_red  = rl_colors.HexColor(_RED_HEX)
_hf_blue = rl_colors.HexColor(_KPMG_BLUE_HEX)
_hf_pac  = rl_colors.HexColor(_PACIFIC_HEX)
_hf_lt   = rl_colors.HexColor(_LIGHT_HEX)

_HF_PS_CONF     = ParagraphStyle("hf_conf",     parent=_hf_base["Normal"], fontSize=9,  textColor=_hf_red,  alignment=TA_RIGHT, fontName="Helvetica-Bold")
_HF_PS_TITLE    = ParagraphStyle("hf_title",    parent=_hf_base["Normal"], fontSize=16, alignment=TA_CENTER, fontName="Helvetica-Bold", spaceAfter=2)
_HF_PS_SUBTITLE = ParagraphStyle("hf_subtitle", parent=_hf_base["Normal"], fontSize=12, alignment=TA_CENTER, spaceAfter=4)
_HF_PS_NORMAL   = ParagraphStyle("hf_normal",   parent=_hf_base["Normal"], fontSize=10, leading=14)
_HF_PS_INDENT   = ParagraphStyle("hf_indent",   parent=_hf_base["Normal"], fontSize=10, leading=14, leftIndent=12)
_HF_PS_VALUE    = ParagraphStyle("hf_value",    parent=_hf_base["Normal"], fontSize=10, leading=14, alignment=TA_RIGHT)
_HF_PS_BOLD     = ParagraphStyle("hf_bold",     parent=_hf_base["Normal"], fontSize=10, leading=14, fontName="Helvetica-Bold")
_HF_PS_BOLDV    = ParagraphStyle("hf_boldv",    parent=_hf_base["Normal"], fontSize=10, leading=14, fontName="Helvetica-Bold", alignment=TA_RIGHT)
_HF_PS_FOOT     = ParagraphStyle("hf_foot",     parent=_hf_base["Normal"], fontSize=8,  fontName="Helvetica-Oblique")
_HF_PS_SEC      = ParagraphStyle("hf_sec",      parent=_hf_base["Normal"], fontSize=10, fontName="Helvetica-Bold", textColor=_hf_blue, spaceBefore=6, spaceAfter=3)
_HF_PS_TH       = ParagraphStyle("hf_th",       parent=_hf_base["Normal"], fontSize=8,  fontName="Helvetica-Bold", alignment=TA_CENTER, textColor=rl_colors.white)
_HF_PS_TH_L     = ParagraphStyle("hf_th_l",     parent=_hf_base["Normal"], fontSize=8,  fontName="Helvetica-Bold", textColor=rl_colors.white)
_HF_PS_TD_L     = ParagraphStyle("hf_td_l",     parent=_hf_base["Normal"], fontSize=9,  leading=12)
_HF_PS_TD_R     = ParagraphStyle("hf_td_r",     parent=_hf_base["Normal"], fontSize=9,  leading=12, alignment=TA_RIGHT)
_HF_PS_TD_LB    = ParagraphStyle("hf_td_lb",    parent=_hf_base["Normal"], fontSize=9,  leading=12, fontName="Helvetica-Bold")
_HF_PS_TD_RB    = ParagraphStyle("hf_td_rb",    parent=_hf_base["Normal"], fontSize=9,  leading=12, fontName="Helvetica-Bold", alignment=TA_RIGHT)


# ══════════════════════════════════════════════════════════════════════════════
# Word (.docx) helpers
# ══════════════════════════════════════════════════════════════════════════════

def _set_run(run, bold=False, underline=False, size_pt=10,
             color: RGBColor | None = None) -> None:
    run.bold = bold
    run.underline = underline
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color


def _two_col(doc: Document, label: str, value: str,
             bold: bool = False, indent: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(4.5), WD_ALIGN_PARAGRAPH.RIGHT)
    prefix = "    " if indent else ""
    run = p.add_run(f"{prefix}{label}\t{value}")
    _set_run(run, bold=bold)


def _section_hdr(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run(run, bold=True, underline=True, size_pt=10, color=_KPMG_BLUE)


def _kv_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    _set_run(r1, bold=True)
    r2 = p.add_run(value)
    _set_run(r2)


def _blank(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


def _shaded_cell(cell, hex_fill: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"),  hex_fill)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"),   "clear")
    tcPr.append(shd)


def _add_table_docx(doc: Document, headers: list[str],
                    data_rows: list[list[str]],
                    bold_last: bool = False) -> None:
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    tbl.style = "Table Grid"

    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold           = True
            run.font.size      = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shaded_cell(hdr[i], "00338D")

    for ri, row_data in enumerate(data_rows):
        is_bold_row = bold_last and (ri == len(data_rows) - 1)
        cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
            align = WD_ALIGN_PARAGRAPH.RIGHT if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            cells[ci].paragraphs[0].alignment = align
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(9)
                run.bold      = is_bold_row
            if is_bold_row:
                _shaded_cell(cells[ci], "E8F4FF")


# ══════════════════════════════════════════════════════════════════════════════
# Word (.docx) builder
# ══════════════════════════════════════════════════════════════════════════════

def _build_hf_docx_from_series(row: pd.Series) -> Document:
    _check_cq_arithmetic(row)  # Change 7: raises ValueError on roll-forward discrepancy > $1
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)

    # ── CONFIDENTIAL + title ──────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("CONFIDENTIAL")
    _set_run(r, bold=True, size_pt=9, color=RGBColor(0xD7, 0x3B, 0x3E))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p = doc.add_paragraph()
    r = p.add_run("Hedge Fund Capital Account Statement")
    _set_run(r, bold=True, size_pt=16, color=_KPMG_BLUE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    r = p.add_run("Current Reporting Period")
    _set_run(r, size_pt=11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _blank(doc)
    _kv_line(doc, "Investor",   _gs(row, "INVESTOR_NAME"))
    _kv_line(doc, "Inception",  _gs(row, "INCEPTION_DATE"))
    _kv_line(doc, "Currency",   _gs(row, "REPT_CCY"))
    _kv_line(doc, "% of AUM",   _fmt_pct(_g(row, "PCT_AUM")))
    _blank(doc)

    # ── 1: Capital Account Summary ────────────────────────────────────────────
    _section_hdr(doc, "1. Capital Account Summary")
    _blank(doc)
    _add_table_docx(doc,
        headers=["", "Current Quarter (CQ)", "Year-to-Date (YTD)", "Inception-to-Date (ITD)"],
        bold_last=True,
        data_rows=[
            ["Beginning Partner's Capital",    fmt_usd(_g(row,"BEG_CAP_CQ")),  fmt_usd(_g(row,"BEG_CAP_YTD")),  fmt_usd(_g(row,"BEG_CAP_ITD"))],
            ["(+) Capital Contributions",      fmt_usd(_g(row,"CONTRIB_CQ")),  fmt_usd(_g(row,"CONTRIB_YTD")),  fmt_usd(_g(row,"CONTRIB_ITD"))],
            ["(+) DRIP Reinvestment",          fmt_usd(_g(row,"DRIP_CQ")),     fmt_usd(_g(row,"DRIP_YTD")),     fmt_usd(_g(row,"DRIP_ITD"))],
            ["(+) Investment Income (Loss)",   fmt_usd(_g(row,"INC_CQ")),      fmt_usd(_g(row,"INC_YTD")),      fmt_usd(_g(row,"INC_ITD"))],
            ["(+) Net Unrealized Gain (Loss)", fmt_usd(_g(row,"UNRLZ_CQ")),    fmt_usd(_g(row,"UNRLZ_YTD")),    fmt_usd(_g(row,"UNRLZ_ITD"))],
            ["(+) Net Realized Gain (Loss)",   fmt_usd(_g(row,"RLZD_CQ")),     fmt_usd(_g(row,"RLZD_YTD")),     fmt_usd(_g(row,"RLZD_ITD"))],
            ["(–) Distributions to LP",        fmt_usd(_g(row,"DIST_LP_CQ")),  fmt_usd(_g(row,"DIST_LP_YTD")),  fmt_usd(_g(row,"DIST_LP_ITD"))],
            ["(–) Management Fees",            fmt_usd(_g(row,"DIST_MGR_CQ")), fmt_usd(_g(row,"DIST_MGR_YTD")), fmt_usd(_g(row,"DIST_MGR_ITD"))],
            ["(–) Incentive Fees",             fmt_usd(_g(row,"INC_FEE_CQ")),  fmt_usd(_g(row,"INC_FEE_YTD")),  fmt_usd(_g(row,"INC_FEE_ITD"))],
            ["Ending Partner's Capital",       fmt_usd(_g(row,"END_CAP_CQ")),  fmt_usd(_g(row,"END_CAP_YTD")),  fmt_usd(_g(row,"END_CAP_ITD"))],
        ],
    )
    _blank(doc)

    # ── 2: Unit Reconciliation ────────────────────────────────────────────────
    _section_hdr(doc, "2. Unit Reconciliation")
    _blank(doc)
    _end_units_itd = _g(row, "END_UNITS_ITD") + _g(row, "DRIP_U_ITD")  # Change 8: include DRIP
    _add_table_docx(doc,
        headers=["", "CQ Units", "YTD Units", "ITD Units"],
        bold_last=True,
        data_rows=[
            ["Beginning Units",    _fmt_units(_g(row,"BEG_UNITS_CQ")),   _fmt_units(_g(row,"BEG_UNITS_YTD")),  _fmt_units(_g(row,"BEG_UNITS_ITD"))],
            ["(+) Units Issued",   _fmt_units(_g(row,"CONTRIB_U_CQ")),   _fmt_units(_g(row,"CONTRIB_U_YTD")),  _fmt_units(_g(row,"CONTRIB_U_ITD"))],
            ["(–) Units Redeemed", _fmt_units(_g(row,"REDEMP_U_CQ")),    _fmt_units(_g(row,"REDEMP_U_YTD")),   _fmt_units(_g(row,"REDEMP_U_ITD"))],
            ["(+) DRIP Units",     _fmt_units(_g(row,"DRIP_U_CQ")),      _fmt_units(_g(row,"DRIP_U_YTD")),     _fmt_units(_g(row,"DRIP_U_ITD"))],
            ["Ending Units",       _fmt_units(_g(row,"END_UNITS_CQ")),   _fmt_units(_g(row,"END_UNITS_YTD")),  _fmt_units(_end_units_itd)],
        ],
    )
    if _end_units_itd != 0.0:
        _two_col(doc, "NAV per Unit (ITD)", _fmt_px(_g(row, "END_CAP_ITD") / _end_units_itd))
    else:
        print(f"WARN: {_gs(row,'INVESTOR_NAME')} — Ending Units ITD is zero after DRIP inclusion; NAV per unit not computed")
    _blank(doc)

    # ── 3: Unit Price Attribution ─────────────────────────────────────────────
    _section_hdr(doc, "3. Unit Price Attribution")
    _blank(doc)
    _add_table_docx(doc,
        headers=["Line Item", "Unit Price"],
        bold_last=True,
        data_rows=[
            ["Beginning Partner's Capital",            _fmt_px(_g(row,"BEG_PX"))],
            ["Transfer In",                            _fmt_px(_g(row,"XFER_IN_PX"))],
            ["Transfer Out",                           _fmt_px(_g(row,"XFER_OUT_PX"))],
            ["Capital Contribution",                   _fmt_px(_g(row,"CONTRIB_PX"))],
            ["Investment Income",                      _fmt_px(_g(row,"INC_PX"))],
            ["Fund Level Expense",                     _fmt_px(_g(row,"EXP_PX"))],
            ["Net Unrealized Gain / (Loss)",           _fmt_px(_g(row,"UNRLZ_PX"))],
            ["Net Realized Gain / (Loss)",             _fmt_px(_g(row,"RLZD_PX"))],
            ["Partner's Equity Before Distributions",  _fmt_px(_g(row,"EQ_PRED_PX"))],
            ["Distributions to LPs",                  _fmt_px(_g(row,"DIST_LP_PX"))],
            ["Distributions to Manager",               _fmt_px(_g(row,"DIST_MGR_PX"))],
            ["Incentive Fee",                          _fmt_px(_g(row,"INC_FEE_PX"))],
            ["Tax Reduction on Distributions",         _fmt_px(_g(row,"TAX_RED_PX"))],
            ["Ending Partner's Capital",               _fmt_px(_g(row,"END_PX"))],
        ],
    )
    _blank(doc)

    # ── 4: Performance Analytics ──────────────────────────────────────────────
    _section_hdr(doc, "4. Performance Analytics")
    _blank(doc)
    _two_col(doc, "Gross IRR",           _fmt_pct(_g(row,"GROSS_IRR")))
    _two_col(doc, "Net IRR",             _fmt_pct(_g(row,"NET_IRR")))
    _two_col(doc, "DPI",                 _fmt_x(_g(row,"DPI")))
    _two_col(doc, "RVPI",                _fmt_x(_g(row,"RVPI")))
    _two_col(doc, "TVPI",                _fmt_x(_g(row,"TVPI")))
    _two_col(doc, "Total Return (CQ)",   _fmt_pct(_g(row,"TOT_RET_CQ_PCT")))
    _two_col(doc, "Total Return (ITD)",  _fmt_pct(_g(row,"TOT_RET_ITD_PCT")))
    _two_col(doc, "Preferred Return",    _fmt_pct(_g(row,"PREF_RET")))
    _two_col(doc, "Incentive Fee Rate",  _fmt_pct(_g(row,"INC_FEE_RATE")))
    _two_col(doc, "Hurdle Exceeded?",    _gs(row,"HURDLE_EXCEEDED"))
    _blank(doc)

    # ── 5: Capital Commitment ─────────────────────────────────────────────────
    _section_hdr(doc, "5. Capital Commitment")
    _blank(doc)
    _two_col(doc, "Total Commitment",     fmt_usd(_g(row,"TOTAL_COMMIT")))
    _two_col(doc, "Funded Commitment",    fmt_usd(_g(row,"FUNDED_COMMIT")))
    _two_col(doc, "Transferred",          fmt_usd(_g(row,"XFER_COMMIT")))
    _two_col(doc, "Available Commitment", fmt_usd(_g(row,"AVAIL_COMMIT")))
    _two_col(doc, "Commitment Funded",    _fmt_pct(_g(row,"COMMIT_FUNDED_PCT")))
    _blank(doc)

    # ── 6: Return & Fee Analysis ──────────────────────────────────────────────
    _section_hdr(doc, "6. Return & Fee Analysis")
    _blank(doc)
    _add_table_docx(doc,
        headers=["Metric", "CQ ($)", "ITD ($)"],
        data_rows=[
            ["Unrealized Gain",  fmt_usd(_g(row,"UNRLZ_CQ_DLR")),  fmt_usd(_g(row,"UNRLZ_ITD_DLR"))],
            ["Realized Gain",    fmt_usd(_g(row,"RLZD_CQ_DLR")),   fmt_usd(_g(row,"RLZD_ITD_DLR"))],
            ["Total Return",     fmt_usd(_g(row,"TOT_RET_CQ_DLR")), fmt_usd(_g(row,"TOT_RET_ITD_DLR"))],
            ["Mgmt. Fee",        "—",                               fmt_usd(_g(row,"MGMT_FEE_ITD_DLR"))],
            ["Incentive Fee",    "—",                               fmt_usd(_g(row,"INC_FEE_ITD_DLR"))],
            ["Total Fees",       "—",                               fmt_usd(_g(row,"TOT_FEES_ITD_DLR"))],
            ["Net Return",       "—",                               fmt_usd(_g(row,"NET_RET_ITD_DLR"))],
        ],
        bold_last=True,
    )
    _blank(doc)

    # ── 7: Waterfall Analysis ─────────────────────────────────────────────────
    _section_hdr(doc, "7. Waterfall Analysis")
    _blank(doc)
    _two_col(doc, "Hurdle Amount (ITD)",     fmt_usd(_g(row,"HURDLE_AMT_ITD")))
    _two_col(doc, "Excess Over Hurdle",      fmt_usd(_g(row,"EXCESS_HURDLE")))
    _two_col(doc, "GP Catch-Up Amount",      fmt_usd(_g(row,"GP_CATCHUP_AMT")))
    _two_col(doc, "LP Net Waterfall Share",  fmt_usd(_g(row,"LP_NET_WF")), bold=True)
    _two_col(doc, "Waterfall Tier",          _gs(row,"WF_TIER"))
    _blank(doc)

    # ── 8: Lock-up & Redemption Terms ────────────────────────────────────────
    _section_hdr(doc, "8. Lock-up & Redemption Terms")
    _blank(doc)
    _two_col(doc, "Lock-up Period",               _fmt_mo(_g(row,"LOCKUP_MO")))
    _two_col(doc, "Subscription Date",            _gs(row,"SUB_DATE"))
    _two_col(doc, "First Contribution Date",      _gs(row,"FIRST_CONTRIB_DATE"))
    _two_col(doc, "Redemption Eligibility Date",  _gs(row,"REDEMP_ELIG_DATE"))
    _two_col(doc, "Lock-up Expired?",             _gs(row,"LOCKUP_EXPIRED"))
    _two_col(doc, "Months Remaining",             _fmt_int(_g(row,"MONTHS_REM"), " mo"))
    _two_col(doc, "Redemption Frequency",         _gs(row,"REDEMP_FREQ"))
    _two_col(doc, "Gate Provision",               _gs(row,"GATE_PROV"))
    _two_col(doc, "Notice Period",                _fmt_int(_g(row,"NOTICE_DAYS"), " days"))
    _two_col(doc, "Side Pocket Eligible?",        _gs(row,"SIDE_POCKET"))
    _two_col(doc, "DRIP Enrolled?",               _gs(row,"DRIP_ENROLLED"))
    _two_col(doc, "Distribution Preference",      _gs(row,"DIST_PREF"))
    _two_col(doc, "High-Water Mark Active?",      _gs(row,"HWM_ACTIVE"))
    _blank(doc)

    # ── 9: Side Letter & Compliance ───────────────────────────────────────────
    _section_hdr(doc, "9. Side Letter & Compliance")
    _blank(doc)
    _two_col(doc, "Side Letter Flag",               _gs(row,"SIDE_LETTER_FLG"))
    _two_col(doc, "Mgmt. Fee Rate (Side Letter)",   _fmt_pct(_g(row,"MGMT_FEE_RATE")))
    _two_col(doc, "Hurdle Type",                    _gs(row,"HURDLE_TYPE"))
    _two_col(doc, "Catch-up %",                     _fmt_pct(_g(row,"CATCHUP_PCT")))
    _two_col(doc, "Reporting Currency",             _gs(row,"REPT_CCY"))
    _two_col(doc, "FATCA Status",                   _gs(row,"FATCA"))
    _two_col(doc, "AML / KYC Status",               _gs(row,"AML_KYC"))
    _two_col(doc, "Accredited / Qualified",         _gs(row,"ACCREDITED"))
    _two_col(doc, "Custodian / Prime Broker",       _gs(row,"CUSTODIAN"))
    _two_col(doc, "Special Terms",                  _gs(row,"SPECIAL_TERMS"))
    _blank(doc)

    # ── Footer ────────────────────────────────────────────────────────────────
    p = doc.add_paragraph(
        "This document is CONFIDENTIAL and prepared for the named investor only. "
        "Capital account balances are based on pre-calculated PCAP data and may not "
        "represent amounts ultimately realized. © KPMG International"
    )
    for run in p.runs:
        run.font.size = Pt(8)

    return doc


# ══════════════════════════════════════════════════════════════════════════════
# PDF builder
# ══════════════════════════════════════════════════════════════════════════════

def build_hf_pdf(row: pd.Series) -> bytes:
    _check_cq_arithmetic(row)  # Change 7: raises ValueError on roll-forward discrepancy > $1
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        topMargin=0.75*inch, bottomMargin=0.75*inch,
        leftMargin=1.0*inch, rightMargin=1.0*inch,
    )

    KB  = rl_colors.HexColor(_KPMG_BLUE_HEX)
    TOT = rl_colors.HexColor(_TOTAL_HEX)
    W   = 6.5 * inch   # usable page width

    # ── PDF table builders ────────────────────────────────────────────────────

    def _kv(label, value):
        return Paragraph(f"<b>{label}:</b> {value}", _HF_PS_NORMAL)

    def _sec(text):
        return Paragraph(f"<b><u>{text}</u></b>", _HF_PS_SEC)

    def _two(label, value, bold=False):
        ls = _HF_PS_BOLD  if bold else _HF_PS_NORMAL
        vs = _HF_PS_BOLDV if bold else _HF_PS_VALUE
        t = Table([[Paragraph(label, ls), Paragraph(value, vs)]],
                  colWidths=[W * 0.62, W * 0.38])
        t.setStyle(TableStyle([
            ("TOPPADDING",    (0,0), (-1,-1), 1),
            ("BOTTOMPADDING", (0,0), (-1,-1), 1),
            ("LEFTPADDING",   (0,0), (-1,-1), 0),
            ("RIGHTPADDING",  (-1,0),(-1,-1), 0),
        ]))
        return t

    def _grid(headers, rows, col_widths, bold_last=False):
        data = [[Paragraph(h, _HF_PS_TH) for h in headers]]
        for ri, rd in enumerate(rows):
            is_bold = bold_last and (ri == len(rows) - 1)
            sl = _HF_PS_TD_LB if is_bold else _HF_PS_TD_L
            sr = _HF_PS_TD_RB if is_bold else _HF_PS_TD_R
            data.append([
                Paragraph(str(v), sr if i > 0 else sl)
                for i, v in enumerate(rd)
            ])
        t = Table(data, colWidths=col_widths)
        ts = [
            ("BACKGROUND",    (0, 0), (-1, 0), KB),
            ("TEXTCOLOR",     (0, 0), (-1, 0), rl_colors.white),
            ("GRID",          (0, 0), (-1, -1), 0.4, _hf_lt),
            ("BACKGROUND",    (0, 1), (-1, -1), rl_colors.white),
            ("TOPPADDING",    (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING",   (0, 0), (-1, -1), 4),
            ("RIGHTPADDING",  (-1,0), (-1, -1), 4),
        ]
        if bold_last and len(rows) > 0:
            lr = len(rows)
            ts.append(("BACKGROUND", (0, lr), (-1, lr), TOT))
        t.setStyle(TableStyle(ts))
        return t

    sp = lambda h=5: Spacer(1, h)
    story = []

    # Header block
    story += [Paragraph("CONFIDENTIAL", _HF_PS_CONF), sp(4)]
    story += [Paragraph("Hedge Fund Capital Account Statement", _HF_PS_TITLE), sp(4)]
    story += [Paragraph("Current Reporting Period", _HF_PS_SUBTITLE), sp(8)]
    story += [_kv("Investor",  _gs(row,"INVESTOR_NAME")), sp(3)]
    story += [_kv("Inception", _gs(row,"INCEPTION_DATE")), sp(3)]
    story += [_kv("Currency",  _gs(row,"REPT_CCY")), sp(3)]
    story += [_kv("% of AUM",  _fmt_pct(_g(row,"PCT_AUM"))), sp(10)]

    # Section 1
    CW4 = [W*0.34, W*0.22, W*0.22, W*0.22]
    story += [_sec("1. Capital Account Summary"), sp(4)]
    story.append(_grid(
        ["", "Current Quarter", "Year-to-Date", "Inception-to-Date"],
        [
            ["Beginning Partner's Capital",    fmt_usd(_g(row,"BEG_CAP_CQ")),  fmt_usd(_g(row,"BEG_CAP_YTD")),  fmt_usd(_g(row,"BEG_CAP_ITD"))],
            ["(+) Capital Contributions",      fmt_usd(_g(row,"CONTRIB_CQ")),  fmt_usd(_g(row,"CONTRIB_YTD")),  fmt_usd(_g(row,"CONTRIB_ITD"))],
            ["(+) DRIP Reinvestment",          fmt_usd(_g(row,"DRIP_CQ")),     fmt_usd(_g(row,"DRIP_YTD")),     fmt_usd(_g(row,"DRIP_ITD"))],
            ["(+) Investment Income (Loss)",   fmt_usd(_g(row,"INC_CQ")),      fmt_usd(_g(row,"INC_YTD")),      fmt_usd(_g(row,"INC_ITD"))],
            ["(+) Net Unrealized Gain (Loss)", fmt_usd(_g(row,"UNRLZ_CQ")),    fmt_usd(_g(row,"UNRLZ_YTD")),    fmt_usd(_g(row,"UNRLZ_ITD"))],
            ["(+) Net Realized Gain (Loss)",   fmt_usd(_g(row,"RLZD_CQ")),     fmt_usd(_g(row,"RLZD_YTD")),     fmt_usd(_g(row,"RLZD_ITD"))],
            ["(–) Distributions to LP",        fmt_usd(_g(row,"DIST_LP_CQ")),  fmt_usd(_g(row,"DIST_LP_YTD")),  fmt_usd(_g(row,"DIST_LP_ITD"))],
            ["(–) Management Fees",            fmt_usd(_g(row,"DIST_MGR_CQ")), fmt_usd(_g(row,"DIST_MGR_YTD")), fmt_usd(_g(row,"DIST_MGR_ITD"))],
            ["(–) Incentive Fees",             fmt_usd(_g(row,"INC_FEE_CQ")),  fmt_usd(_g(row,"INC_FEE_YTD")),  fmt_usd(_g(row,"INC_FEE_ITD"))],
            ["Ending Partner's Capital",       fmt_usd(_g(row,"END_CAP_CQ")),  fmt_usd(_g(row,"END_CAP_YTD")),  fmt_usd(_g(row,"END_CAP_ITD"))],
        ],
        col_widths=CW4, bold_last=True,
    ))
    story.append(sp(10))

    # Section 2
    _end_units_itd = _g(row, "END_UNITS_ITD") + _g(row, "DRIP_U_ITD")  # Change 8: include DRIP
    story += [_sec("2. Unit Reconciliation"), sp(4)]
    story.append(_grid(
        ["", "CQ Units", "YTD Units", "ITD Units"],
        [
            ["Beginning Units",    _fmt_units(_g(row,"BEG_UNITS_CQ")),   _fmt_units(_g(row,"BEG_UNITS_YTD")),  _fmt_units(_g(row,"BEG_UNITS_ITD"))],
            ["(+) Units Issued",   _fmt_units(_g(row,"CONTRIB_U_CQ")),   _fmt_units(_g(row,"CONTRIB_U_YTD")),  _fmt_units(_g(row,"CONTRIB_U_ITD"))],
            ["(–) Units Redeemed", _fmt_units(_g(row,"REDEMP_U_CQ")),    _fmt_units(_g(row,"REDEMP_U_YTD")),   _fmt_units(_g(row,"REDEMP_U_ITD"))],
            ["(+) DRIP Units",     _fmt_units(_g(row,"DRIP_U_CQ")),      _fmt_units(_g(row,"DRIP_U_YTD")),     _fmt_units(_g(row,"DRIP_U_ITD"))],
            ["Ending Units",       _fmt_units(_g(row,"END_UNITS_CQ")),   _fmt_units(_g(row,"END_UNITS_YTD")),  _fmt_units(_end_units_itd)],
        ],
        col_widths=CW4, bold_last=True,
    ))
    if _end_units_itd != 0.0:
        story += [sp(4), _two("NAV per Unit (ITD)", _fmt_px(_g(row, "END_CAP_ITD") / _end_units_itd))]
    else:
        print(f"WARN: {_gs(row,'INVESTOR_NAME')} — Ending Units ITD is zero after DRIP inclusion; NAV per unit not computed")
    story.append(sp(10))

    # Section 3
    CW2A = [W * 0.55, W * 0.45]
    story += [_sec("3. Unit Price Attribution"), sp(4)]
    story.append(_grid(
        ["Line Item", "Unit Price"],
        [
            ["Beginning Partner's Capital",           _fmt_px(_g(row,"BEG_PX"))],
            ["Transfer In",                           _fmt_px(_g(row,"XFER_IN_PX"))],
            ["Transfer Out",                          _fmt_px(_g(row,"XFER_OUT_PX"))],
            ["Capital Contribution",                  _fmt_px(_g(row,"CONTRIB_PX"))],
            ["Investment Income",                     _fmt_px(_g(row,"INC_PX"))],
            ["Fund Level Expense",                    _fmt_px(_g(row,"EXP_PX"))],
            ["Net Unrealized Gain / (Loss)",          _fmt_px(_g(row,"UNRLZ_PX"))],
            ["Net Realized Gain / (Loss)",            _fmt_px(_g(row,"RLZD_PX"))],
            ["Partner's Equity Before Distributions", _fmt_px(_g(row,"EQ_PRED_PX"))],
            ["Distributions to LPs",                 _fmt_px(_g(row,"DIST_LP_PX"))],
            ["Distributions to Manager",              _fmt_px(_g(row,"DIST_MGR_PX"))],
            ["Incentive Fee",                         _fmt_px(_g(row,"INC_FEE_PX"))],
            ["Tax Reduction on Distributions",        _fmt_px(_g(row,"TAX_RED_PX"))],
            ["Ending Partner's Capital",              _fmt_px(_g(row,"END_PX"))],
        ],
        col_widths=CW2A, bold_last=True,
    ))
    story.append(sp(10))

    # Section 4
    story += [_sec("4. Performance Analytics"), sp(4)]
    for lbl, val in [
        ("Gross IRR",           _fmt_pct(_g(row,"GROSS_IRR"))),
        ("Net IRR",             _fmt_pct(_g(row,"NET_IRR"))),
        ("DPI",                 _fmt_x(_g(row,"DPI"))),
        ("RVPI",                _fmt_x(_g(row,"RVPI"))),
        ("TVPI",                _fmt_x(_g(row,"TVPI"))),
        ("Total Return (CQ)",   _fmt_pct(_g(row,"TOT_RET_CQ_PCT"))),
        ("Total Return (ITD)",  _fmt_pct(_g(row,"TOT_RET_ITD_PCT"))),
        ("Preferred Return",    _fmt_pct(_g(row,"PREF_RET"))),
        ("Incentive Fee Rate",  _fmt_pct(_g(row,"INC_FEE_RATE"))),
        ("Hurdle Exceeded?",    _gs(row,"HURDLE_EXCEEDED")),
    ]:
        story.append(_two(lbl, val))
    story.append(sp(10))

    # Section 5
    story += [_sec("5. Capital Commitment"), sp(4)]
    for lbl, val in [
        ("Total Commitment",     fmt_usd(_g(row,"TOTAL_COMMIT"))),
        ("Funded Commitment",    fmt_usd(_g(row,"FUNDED_COMMIT"))),
        ("Transferred",          fmt_usd(_g(row,"XFER_COMMIT"))),
        ("Available Commitment", fmt_usd(_g(row,"AVAIL_COMMIT"))),
        ("Commitment Funded",    _fmt_pct(_g(row,"COMMIT_FUNDED_PCT"))),
    ]:
        story.append(_two(lbl, val))
    story.append(sp(10))

    # Section 6
    CW3 = [W*0.40, W*0.30, W*0.30]
    story += [_sec("6. Return & Fee Analysis"), sp(4)]
    story.append(_grid(
        ["Metric", "CQ ($)", "ITD ($)"],
        [
            ["Unrealized Gain",  fmt_usd(_g(row,"UNRLZ_CQ_DLR")),   fmt_usd(_g(row,"UNRLZ_ITD_DLR"))],
            ["Realized Gain",    fmt_usd(_g(row,"RLZD_CQ_DLR")),    fmt_usd(_g(row,"RLZD_ITD_DLR"))],
            ["Total Return",     fmt_usd(_g(row,"TOT_RET_CQ_DLR")), fmt_usd(_g(row,"TOT_RET_ITD_DLR"))],
            ["Mgmt. Fee",        "—",                                fmt_usd(_g(row,"MGMT_FEE_ITD_DLR"))],
            ["Incentive Fee",    "—",                                fmt_usd(_g(row,"INC_FEE_ITD_DLR"))],
            ["Total Fees",       "—",                                fmt_usd(_g(row,"TOT_FEES_ITD_DLR"))],
            ["Net Return",       "—",                                fmt_usd(_g(row,"NET_RET_ITD_DLR"))],
        ],
        col_widths=CW3, bold_last=True,
    ))
    story.append(sp(10))

    # Section 7
    story += [_sec("7. Waterfall Analysis"), sp(4)]
    for lbl, val, bold in [
        ("Hurdle Amount (ITD)",     fmt_usd(_g(row,"HURDLE_AMT_ITD")), False),
        ("Excess Over Hurdle",      fmt_usd(_g(row,"EXCESS_HURDLE")),  False),
        ("GP Catch-Up Amount",      fmt_usd(_g(row,"GP_CATCHUP_AMT")), False),
        ("LP Net Waterfall Share",  fmt_usd(_g(row,"LP_NET_WF")),      True),
        ("Waterfall Tier",          _gs(row,"WF_TIER"),                False),
    ]:
        story.append(_two(lbl, val, bold=bold))
    story.append(sp(10))

    # Section 8
    story += [_sec("8. Lock-up & Redemption Terms"), sp(4)]
    for lbl, val in [
        ("Lock-up Period",              _fmt_mo(_g(row,"LOCKUP_MO"))),
        ("Subscription Date",           _gs(row,"SUB_DATE")),
        ("First Contribution Date",     _gs(row,"FIRST_CONTRIB_DATE")),
        ("Redemption Eligibility Date", _gs(row,"REDEMP_ELIG_DATE")),
        ("Lock-up Expired?",            _gs(row,"LOCKUP_EXPIRED")),
        ("Months Remaining",            _fmt_int(_g(row,"MONTHS_REM"), " mo")),
        ("Redemption Frequency",        _gs(row,"REDEMP_FREQ")),
        ("Gate Provision",              _gs(row,"GATE_PROV")),
        ("Notice Period",               _fmt_int(_g(row,"NOTICE_DAYS"), " days")),
        ("Side Pocket Eligible?",       _gs(row,"SIDE_POCKET")),
        ("DRIP Enrolled?",              _gs(row,"DRIP_ENROLLED")),
        ("Distribution Preference",     _gs(row,"DIST_PREF")),
        ("High-Water Mark Active?",     _gs(row,"HWM_ACTIVE")),
    ]:
        story.append(_two(lbl, val))
    story.append(sp(10))

    # Section 9
    story += [_sec("9. Side Letter & Compliance"), sp(4)]
    for lbl, val in [
        ("Side Letter Flag",              _gs(row,"SIDE_LETTER_FLG")),
        ("Mgmt. Fee Rate (Side Letter)",  _fmt_pct(_g(row,"MGMT_FEE_RATE"))),
        ("Hurdle Type",                   _gs(row,"HURDLE_TYPE")),
        ("Catch-up %",                    _fmt_pct(_g(row,"CATCHUP_PCT"))),
        ("Reporting Currency",            _gs(row,"REPT_CCY")),
        ("FATCA Status",                  _gs(row,"FATCA")),
        ("AML / KYC Status",              _gs(row,"AML_KYC")),
        ("Accredited / Qualified",        _gs(row,"ACCREDITED")),
        ("Custodian / Prime Broker",      _gs(row,"CUSTODIAN")),
        ("Special Terms",                 _gs(row,"SPECIAL_TERMS")),
    ]:
        story.append(_two(lbl, val))
    story.append(sp(16))

    story.append(Paragraph(
        "This statement is CONFIDENTIAL and prepared for the named investor only. "
        "Balances are based on pre-calculated PCAP data and may not represent amounts "
        "ultimately realized. © KPMG International",
        _HF_PS_FOOT,
    ))

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════════════════════
# Meridian pipeline — dict-based docx builder (11 sections, A4, full layout)
# ══════════════════════════════════════════════════════════════════════════════

from docx.shared import Cm
from datetime import datetime as _dt

_M_BLUE_HEX  = "00338D"
_M_LBL_HEX   = "E8EDF5"
_M_ALT_HEX   = "F0F4FA"
_M_TOTAL_HEX = "DDEEFF"
_M_PAGE_W    = 17.0   # usable width cm (A4 − 2×2 cm margins)


def _m_g(inv: dict, key: str, dft: float = 0.0) -> float:
    v = inv.get(key, dft)
    if v is None or (isinstance(v, float) and pd.isna(v)):
        return dft
    try:
        return float(v)
    except (TypeError, ValueError):
        return dft


def _m_gs(inv: dict, key: str, dft: str = "—") -> str:
    v = inv.get(key, dft)
    s = str(v).strip() if v is not None else ""
    return s if s and s not in ("nan", "None", "") else dft


def _m_fmt_usd(v) -> str:
    try:
        f = float(v)
        return f"(${abs(f):,.2f})" if f < 0 else f"${f:,.2f}"
    except (TypeError, ValueError):
        return "—"


def _m_fmt_pct(v) -> str:
    try:
        return f"{float(v):.2f}%"
    except (TypeError, ValueError):
        return "—"


def _m_fmt_x(v) -> str:
    try:
        f = float(v)
        return f"{f:.2f}x" if f else "—"
    except (TypeError, ValueError):
        return "—"


def _m_fmt_date(s) -> str:
    if not s or str(s).strip() in ("—", "", "nan", "None"):
        return "—"
    for fmt in ("%d-%b-%Y", "%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y", "%b %d, %Y"):
        try:
            return _dt.strptime(str(s).strip(), fmt).strftime("%-d %B %Y")
        except ValueError:
            pass
    return str(s).strip()


def _m_fmt_units(v) -> str:
    try:
        f = float(v)
        return f"{f:,.6f}" if f else "—"
    except (TypeError, ValueError):
        return "—"


def _m_calibri(run, size: int = 9, bold: bool = False, color: RGBColor | None = None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold      = bold
    if color:
        run.font.color.rgb = color


def _m_shd(cell, hex_col: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"),  hex_col)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"),   "clear")
    tcPr.append(shd)


def _m_cell(cell, text: str, bold: bool = False, align: str = "left",
            fill: str | None = None, sz: int = 9,
            color: RGBColor | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(str(text) if text is not None else "")
    _m_calibri(run, sz, bold, color)
    if fill:
        _m_shd(cell, fill)
    # Compact vertical padding
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "60")
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def _m_no_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "none")
        tblBorders.append(elem)
    tblPr.append(tblBorders)


def _m_section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    _m_calibri(run, 11, bold=True, color=_KPMG_BLUE)


def _m_kv_table(doc: Document, rows: list[tuple[str, str]]):
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(6.5)
    tbl.columns[1].width = Cm(10.5)
    for i, (lbl, val) in enumerate(rows):
        alt = _M_ALT_HEX if i % 2 else "FFFFFF"
        _m_cell(tbl.cell(i, 0), lbl, bold=True, fill=_M_LBL_HEX)
        _m_cell(tbl.cell(i, 1), val, fill=alt)


def _m_data_table(doc: Document, headers: list[str], rows: list[list],
                  bold_last: bool = False, col_widths: list[float] | None = None):
    n   = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n)
    tbl.style = "Table Grid"
    if col_widths:
        for j, w in enumerate(col_widths):
            if j < n:
                tbl.columns[j].width = Cm(w)
    # Header row
    W = RGBColor(0xFF, 0xFF, 0xFF)
    for j, h in enumerate(headers):
        _m_cell(tbl.cell(0, j), h, bold=True, align="center",
                fill=_M_BLUE_HEX, color=W)
    # Data rows
    for i, row_data in enumerate(rows):
        is_bold  = bold_last and i == len(rows) - 1
        fill_hex = _M_TOTAL_HEX if is_bold else (_M_ALT_HEX if i % 2 else "FFFFFF")
        for j, val in enumerate(row_data):
            align = "left" if j == 0 else "right"
            _m_cell(tbl.cell(i + 1, j), str(val) if val is not None else "—",
                    bold=is_bold, fill=fill_hex, align=align)


def _m_page_number(paragraph):
    run = paragraph.add_run("Page ")
    _m_calibri(run, 8)
    for instr in (" PAGE ", " NUMPAGES "):
        fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
        run._r.append(fc1)
        it  = OxmlElement("w:instrText"); it.text = instr
        run._r.append(it)
        fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "end")
        run._r.append(fc2)
        if instr == " PAGE ":
            r2 = paragraph.add_run(" of "); _m_calibri(r2, 8)


def _build_meridian_docx(investor: dict, output_path: str) -> str:
    """
    Build a Meridian Capital Account Statement (.docx) from an investor dict
    (as returned by load_pcap).  Saves to output_path and returns it.
    11 sections + header/footer.  A4, Calibri, KPMG colour scheme.
    """
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)

    RED = RGBColor(0xD7, 0x3B, 0x3E)
    inv_name = _m_gs(investor, "INVESTOR_NAME")

    # ── Running header ────────────────────────────────────────────────────────
    hdr = sec.header
    hdr.is_linked_to_previous = False
    for p in hdr.paragraphs:
        p.clear()
    htbl = hdr.add_table(1, 2, Cm(_M_PAGE_W))
    _m_cell(htbl.cell(0, 0), "Meridian Opportunities Fund, L.P.",
            bold=True, color=_KPMG_BLUE)
    p_r = htbl.cell(0, 1).paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_r.add_run("PRIVATE & CONFIDENTIAL")
    _m_calibri(run, 9, bold=True, color=RED)
    _m_no_borders(htbl)

    # ── Running footer ────────────────────────────────────────────────────────
    ftr = sec.footer
    ftr.is_linked_to_previous = False
    for p in ftr.paragraphs:
        p.clear()
    ftbl = ftr.add_table(1, 3, Cm(_M_PAGE_W))
    _m_cell(ftbl.cell(0, 0), f"CONFIDENTIAL — {inv_name}", sz=8)
    p_c = ftbl.cell(0, 1).paragraphs[0]
    p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_c = p_c.add_run("Meridian Capital Management LLC  |  Q1 2026")
    _m_calibri(run_c, 8)
    p_pg = ftbl.cell(0, 2).paragraphs[0]
    p_pg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _m_page_number(p_pg)
    _m_no_borders(ftbl)

    # ── Title block ───────────────────────────────────────────────────────────
    for text, sz, bold in [
        ("Capital Account Statement",                      16, True),
        ("Q1 2026  |  1 January 2026 – 31 March 2026",    12, False),
        ("Prepared by Meridian Capital Management LLC",    11, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        _m_calibri(run, sz, bold, color=_KPMG_BLUE if bold else None)

    # Thin rule
    p_rule = doc.add_paragraph()
    pPr = p_rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), _M_BLUE_HEX)
    pBdr.append(bot)
    pPr.append(pBdr)
    doc.add_paragraph()

    # ── Section 1: Investor Summary ───────────────────────────────────────────
    _m_section_heading(doc, "1.  Investor Summary")
    aml = _m_gs(investor, "AML_KYC")
    _m_kv_table(doc, [
        ("Legal Name",               inv_name),
        ("Entity Type",              _m_gs(investor, "ENTITY_TYPE")),
        ("Tax ID / EIN",             _m_gs(investor, "TAX_ID")),
        ("Jurisdiction",             _m_gs(investor, "JURISDICTION")),
        ("Reporting Currency",       _m_gs(investor, "REPT_CCY", "USD")),
        ("Inception Date",           _m_fmt_date(_m_gs(investor, "INCEPTION_DATE"))),
        ("Custodian / Prime Broker", _m_gs(investor, "CUSTODIAN")),
        ("AML / KYC Status",         aml),
        ("Investor Classification",  _m_gs(investor, "ACCREDITED")),
        ("FATCA Status",             _m_gs(investor, "FATCA")),
    ])
    if aml == "In Review":
        p = doc.add_paragraph()
        run = p.add_run(
            "⚠  AML/KYC IN REVIEW — Capital deployment suspended "
            "pending compliance clearance."
        )
        _m_calibri(run, 10, bold=True, color=RED)
    doc.add_paragraph()

    # ── Section 2: Unit Price Summary ─────────────────────────────────────────
    _m_section_heading(doc, "2.  Unit Price Summary")
    _m_data_table(doc,
        headers=["Beginning", "Transfer In", "Contribution", "Income",
                 "Expense", "Unrealized", "Realized", "Ending"],
        rows=[[
            _m_fmt_usd(_m_g(investor, "BEG_PX")),
            _m_fmt_usd(_m_g(investor, "XFER_IN_PX")),
            _m_fmt_usd(_m_g(investor, "CONTRIB_PX")),
            _m_fmt_usd(_m_g(investor, "INC_PX")),
            _m_fmt_usd(_m_g(investor, "EXP_PX")),
            _m_fmt_usd(_m_g(investor, "UNRLZ_PX")),
            _m_fmt_usd(_m_g(investor, "RLZD_PX")),
            _m_fmt_usd(_m_g(investor, "END_PX")),
        ]],
        col_widths=[2.1]*8,
    )
    doc.add_paragraph()

    # ── Section 3: Capital Account Summary ────────────────────────────────────
    _m_section_heading(doc, "3.  Capital Account Summary")
    def _cap_row(label, cq_k, ytd_k, itd_k):
        return [label,
                _m_fmt_usd(_m_g(investor, cq_k)),
                _m_fmt_usd(_m_g(investor, ytd_k)),
                _m_fmt_usd(_m_g(investor, itd_k))]

    _m_data_table(doc,
        headers=["Line Item", "CQ ($)", "YTD ($)", "ITD ($)"],
        bold_last=True,
        col_widths=[7.0, 3.3, 3.3, 3.3],
        rows=[
            _cap_row("Beginning Capital",                  "BEG_CAP_CQ",  "BEG_CAP_YTD",  "BEG_CAP_ITD"),
            _cap_row("(+) Capital Contributions",          "CONTRIB_CQ",  "CONTRIB_YTD",  "CONTRIB_ITD"),
            _cap_row("(+) DRIP Reinvestment",              "DRIP_CQ",     "DRIP_YTD",     "DRIP_ITD"),
            _cap_row("(−) Capital Redemptions",            "REDEMP_CQ",   "REDEMP_YTD",   "REDEMP_ITD"),
            _cap_row("(+) Transfer In",                    "XFER_IN_CQ",  "XFER_IN_YTD",  "XFER_IN_ITD"),
            _cap_row("(−) Transfer Out",                   "XFER_OUT_CQ", "XFER_OUT_YTD", "XFER_OUT_ITD"),
            _cap_row("(+) Investment Income",              "INC_CQ",      "INC_YTD",      "INC_ITD"),
            _cap_row("(−) Fund Level Expenses",            "EXP_CQ",      "EXP_YTD",      "EXP_ITD"),
            _cap_row("(+) Net Unrealized Gain/(Loss)",     "UNRLZ_CQ",    "UNRLZ_YTD",    "UNRLZ_ITD"),
            _cap_row("(+) Net Realized Gain/(Loss)",       "RLZD_CQ",     "RLZD_YTD",     "RLZD_ITD"),
            _cap_row("Partner's Equity Before Distrib.",   "EQ_PRED_CQ",  "EQ_PRED_YTD",  "EQ_PRED_ITD"),
            _cap_row("(−) Distributions to LP",           "DIST_LP_CQ",  "DIST_LP_YTD",  "DIST_LP_ITD"),
            _cap_row("(−) Fees Redirected to Manager",    "DIST_MGR_CQ", "DIST_MGR_YTD", "DIST_MGR_ITD"),
            _cap_row("(−) Incentive Fee",                  "INC_FEE_CQ",  "INC_FEE_YTD",  "INC_FEE_ITD"),
            _cap_row("Ending Partner's Capital",           "END_CAP_CQ",  "END_CAP_YTD",  "END_CAP_ITD"),
        ],
    )
    doc.add_paragraph()

    # ── Section 4: Units Summary ───────────────────────────────────────────────
    _m_section_heading(doc, "4.  Units Summary")
    def _u_row(label, cq_k, ytd_k, itd_k):
        return [label,
                _m_fmt_units(_m_g(investor, cq_k)),
                _m_fmt_units(_m_g(investor, ytd_k)),
                _m_fmt_units(_m_g(investor, itd_k))]

    _m_data_table(doc,
        headers=["Line Item", "CQ", "YTD", "ITD"],
        bold_last=True,
        col_widths=[7.0, 3.3, 3.3, 3.3],
        rows=[
            _u_row("Beginning Units",   "BEG_UNITS_CQ",  "BEG_UNITS_YTD",  "BEG_UNITS_ITD"),
            _u_row("Units Contributed", "CONTRIB_U_CQ",  "CONTRIB_U_YTD",  "CONTRIB_U_ITD"),
            _u_row("Units from DRIP",   "DRIP_U_CQ",     "DRIP_U_YTD",     "DRIP_U_ITD"),
            _u_row("Units Redeemed",    "REDEMP_U_CQ",   "REDEMP_U_YTD",   "REDEMP_U_ITD"),
            _u_row("Units Transferred In",  "XFER_U_IN_CQ",  "XFER_U_IN_YTD",  "XFER_U_IN_ITD"),
            _u_row("Units Transferred Out", "XFER_U_OUT_CQ", "XFER_U_OUT_YTD", "XFER_U_OUT_ITD"),
            _u_row("Ending Units",      "END_UNITS_CQ",  "END_UNITS_YTD",  "END_UNITS_ITD"),
        ],
    )
    doc.add_paragraph()

    # ── Section 5: Capital Commitment Summary ─────────────────────────────────
    _m_section_heading(doc, "5.  Capital Commitment Summary")
    _m_kv_table(doc, [
        ("Total Capital Commitment",  _m_fmt_usd(_m_g(investor, "TOTAL_COMMIT"))),
        ("Funded Commitment",         _m_fmt_usd(_m_g(investor, "FUNDED_COMMIT"))),
        ("Transfer of Commitment",    _m_fmt_usd(_m_g(investor, "XFER_COMMIT"))),
        ("Available Commitment",      _m_fmt_usd(_m_g(investor, "AVAIL_COMMIT"))),
        ("Commitment Funded %",       _m_fmt_pct(_m_g(investor, "COMMIT_FUNDED_PCT"))),
    ])
    doc.add_paragraph()

    # ── Section 6: Performance Summary ───────────────────────────────────────
    _m_section_heading(doc, "6.  Performance Summary")
    _m_kv_table(doc, [
        ("Gross IRR",            _m_fmt_pct(_m_g(investor, "GROSS_IRR"))),
        ("Net IRR",              _m_fmt_pct(_m_g(investor, "NET_IRR"))),
        ("DPI",                  _m_fmt_x(_m_g(investor, "DPI"))),
        ("RVPI",                 _m_fmt_x(_m_g(investor, "RVPI"))),
        ("TVPI",                 _m_fmt_x(_m_g(investor, "TVPI"))),
        ("% of Fund AUM",        _m_fmt_pct(_m_g(investor, "PCT_AUM"))),
        ("Total Return CQ ($)",  _m_fmt_usd(_m_g(investor, "TOT_RET_CQ_DLR"))),
        ("Total Return CQ %",    _m_fmt_pct(_m_g(investor, "TOT_RET_CQ_PCT"))),
        ("Total Return ITD ($)", _m_fmt_usd(_m_g(investor, "TOT_RET_ITD_DLR"))),
        ("Total Return ITD %",   _m_fmt_pct(_m_g(investor, "TOT_RET_ITD_PCT"))),
    ])
    doc.add_paragraph()

    # ── Section 7: Fee Summary ────────────────────────────────────────────────
    _m_section_heading(doc, "7.  Fee Summary")
    _m_kv_table(doc, [
        ("Management Fee Rate",   _m_fmt_pct(_m_g(investor, "MGMT_FEE_RATE")) + " p.a."),
        ("Incentive Fee Rate",    _m_fmt_pct(_m_g(investor, "INC_FEE_RATE"))),
        ("Preferred Return",      _m_fmt_pct(_m_g(investor, "PREF_RET")) + " p.a."),
        ("Hurdle Type",           _m_gs(investor, "HURDLE_TYPE")),
        ("High-Water Mark Active",_m_gs(investor, "HWM_ACTIVE")),
        ("GP Catch-Up %",         _m_fmt_pct(_m_g(investor, "CATCHUP_PCT"))),
        ("Mgmt Fee ITD",          _m_fmt_usd(_m_g(investor, "MGMT_FEE_ITD_DLR"))),
        ("Incentive Fee ITD",     _m_fmt_usd(_m_g(investor, "INC_FEE_ITD_DLR"))),
        ("Total Fees ITD",        _m_fmt_usd(_m_g(investor, "TOT_FEES_ITD_DLR"))),
        ("Net Return ITD",        _m_fmt_usd(_m_g(investor, "NET_RET_ITD_DLR"))),
    ])
    doc.add_paragraph()

    # ── Section 8: Waterfall Distribution Analysis ────────────────────────────
    _m_section_heading(doc, "8.  Waterfall Distribution Analysis")
    hurdle_r  = _m_g(investor, "PREF_RET", _m_g(investor, "WF_HURDLE_RATE", 8.0))
    gross_pnl = _m_g(investor, "WF_GROSS_PNL", _m_g(investor, "TOT_RET_ITD_DLR"))
    mgmt_fee  = _m_g(investor, "WF_MGMT_FEE",  _m_g(investor, "MGMT_FEE_ITD_DLR"))
    net_pnl   = _m_g(investor, "WF_NET_PNL",   gross_pnl - mgmt_fee)
    hurdle_a  = _m_g(investor, "WF_HURDLE_AMT",_m_g(investor, "HURDLE_AMT_ITD"))
    hurdle_cross = "YES ✓" if net_pnl > hurdle_a else "NO ✗"
    lp_pref   = _m_g(investor, "WF_LP_PREF",   hurdle_a)
    gp_catch  = _m_g(investor, "WF_GP_CATCHUP",_m_g(investor, "GP_CATCHUP_AMT"))
    lp_carry  = _m_g(investor, "WF_LP_CARRY",  _m_g(investor, "EXCESS_HURDLE") * 0.80)
    lp_net    = _m_g(investor, "WF_LP_NET",     _m_g(investor, "LP_NET_WF"))
    end_cap   = _m_g(investor, "WF_END_CAP",    _m_g(investor, "END_CAP_ITD"))
    wf_tier   = _m_gs(investor, "WF_TIER")

    _m_data_table(doc,
        headers=["Step", "Line Item", "Amount ($)"],
        bold_last=False,
        col_widths=[1.5, 10.5, 5.0],
        rows=[
            ["1", "Gross P&L (Before Fees)",                               _m_fmt_usd(gross_pnl)],
            ["1", "Management Fee (Cost)",                                  _m_fmt_usd(-mgmt_fee)],
            ["1", "Net P&L (After Mgmt Fee)",                              _m_fmt_usd(net_pnl)],
            ["2", f"Hurdle Amount ({hurdle_r:.1f}% p.a. / {hurdle_r/4:.2f}% qly)", _m_fmt_usd(hurdle_a)],
            ["2", "Hurdle Crossed?",                                        hurdle_cross],
            ["2", "LP Preferred Return",                                    _m_fmt_usd(lp_pref)],
            ["3", "GP Catch-Up Amount",                                     _m_fmt_usd(gp_catch)],
            ["4", "LP Carry Share",                                         _m_fmt_usd(lp_carry)],
            ["—", "LP Net Allocation (bold)",                               _m_fmt_usd(lp_net)],
            ["—", "Ending Partner Capital",                                 _m_fmt_usd(end_cap)],
            ["—", "Waterfall Tier",                                         wf_tier],
        ],
    )
    doc.add_paragraph()

    # ── Section 9: Lock-Up & Redemption Terms ─────────────────────────────────
    _m_section_heading(doc, "9.  Lock-Up & Redemption Terms")
    _m_kv_table(doc, [
        ("Subscription Date",          _m_fmt_date(_m_gs(investor, "SUB_DATE"))),
        ("First Contribution Date",    _m_fmt_date(_m_gs(investor, "FIRST_CONTRIB_DATE"))),
        ("Redemption Eligibility Date",_m_fmt_date(_m_gs(investor, "REDEMP_ELIG_DATE"))),
        ("Lock-up Period",             f"{int(_m_g(investor, 'LOCKUP_MO', 0))} months"),
        ("Lock-up Expired?",           _m_gs(investor, "LOCKUP_EXPIRED")),
        ("Redemption Frequency",       _m_gs(investor, "REDEMP_FREQ")),
        ("Gate Provision",             _m_gs(investor, "GATE_PROV")),
        ("Notice Period",              f"{int(_m_g(investor, 'NOTICE_DAYS', 0))} days"),
        ("Months Remaining",           str(int(_m_g(investor, "MONTHS_REM", 0)))),
    ])
    doc.add_paragraph()

    # ── Section 10: Side Letter & Compliance ──────────────────────────────────
    _m_section_heading(doc, "10.  Side Letter & Compliance")
    _m_kv_table(doc, [
        ("Side Pocket Eligible",    _m_gs(investor, "SIDE_POCKET")),
        ("DRIP Enrolled",           _m_gs(investor, "DRIP_ENROLLED")),
        ("Distribution Preference", _m_gs(investor, "DIST_PREF")),
        ("Side Letter Flag",        _m_gs(investor, "SIDE_LETTER_FLG")),
        ("Special Terms",           _m_gs(investor, "SPECIAL_TERMS")),
    ])
    doc.add_paragraph()

    # ── Section 11: Transaction History ───────────────────────────────────────
    txns = investor.get("transactions", [])
    if txns:
        _m_section_heading(doc, "11.  Transaction History")
        txn_rows = sorted(
            [t for t in txns if t.get("date")],
            key=lambda t: str(t.get("date", "")),
        )
        data_rows = []
        for t in txn_rows:
            amt = t.get("amount")
            try:
                amt_s = _m_fmt_usd(float(amt))
            except (TypeError, ValueError):
                amt_s = "—"
            data_rows.append([
                str(t.get("txn_id", "—")),
                str(t.get("date",   "—")),
                str(t.get("quarter","—")),
                str(t.get("type",   "—")),
                str(t.get("sub_type","—")),
                amt_s,
                _m_fmt_units(t.get("units")),
                _m_fmt_usd(t.get("unit_price")),
                str(t.get("status", "—")),
            ])
        _m_data_table(doc,
            headers=["Txn ID", "Date", "Quarter", "Type", "Sub-Type",
                     "Amount ($)", "Units", "Unit Price", "Status"],
            rows=data_rows,
            col_widths=[2.0, 2.3, 2.0, 2.2, 2.0, 2.5, 1.8, 2.0, 2.2],
        )

    import os as _os
    _os.makedirs(_os.path.dirname(_os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    return output_path


def build_hf_docx(row_or_investor, output_path: str | None = None):
    """
    Dispatch wrapper — backward-compatible.

    - pd.Series  (from api.py)     → returns Document object (legacy path)
    - dict       (from load_pcap)  → saves to output_path, returns str
    """
    if isinstance(row_or_investor, dict):
        if output_path is None:
            raise ValueError("output_path is required when investor is a dict")
        return _build_meridian_docx(row_or_investor, output_path)
    return _build_hf_docx_from_series(row_or_investor)
